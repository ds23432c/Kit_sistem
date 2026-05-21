from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
from django.utils import timezone
from .models import RepairRequest, ChangeRequest
from core.models import Room, Notification
from accounts.models import User
import io


def send_notification(user, notif_type, title, message, link=''):
    Notification.objects.create(user=user, type=notif_type, title=title, message=message, link=link)


@login_required
def request_list(request):
    qs = RepairRequest.objects.select_related('room', 'created_by', 'assigned_to').all()
    if request.user.is_keeper:
        qs = qs.filter(created_by=request.user)
    elif request.user.is_technician:
        qs = qs.filter(assigned_to=request.user)
    status = request.GET.get('status', '')
    if status:
        qs = qs.filter(status=status)
    return render(request, 'requests_app/list.html', {
        'requests': qs,
        'statuses': RepairRequest.STATUS_CHOICES,
        'current_status': status,
    })


@login_required
def request_detail(request, pk):
    req = get_object_or_404(RepairRequest, pk=pk)
    technicians = User.objects.filter(role=User.ROLE_TECHNICIAN)
    return render(request, 'requests_app/detail.html', {'req': req, 'technicians': technicians})


@login_required
def request_create(request):
    rooms = request.user.rooms_kept.all() if request.user.is_keeper else Room.objects.all()
    from equipment.models import Equipment
    if request.method == 'POST':
        room_id = request.POST.get('room')
        equipment_id = request.POST.get('equipment')
        description = request.POST.get('description', '')
        priority = request.POST.get('priority', 'medium')
        photo = request.FILES.get('photo')
        room = get_object_or_404(Room, pk=room_id)
        equipment = Equipment.objects.filter(pk=equipment_id).first() if equipment_id else None
        req = RepairRequest.objects.create(
            room=room, equipment=equipment, description=description,
            priority=priority, photo=photo, created_by=request.user, status='new'
        )
        # Notify AHC
        for ahc_user in User.objects.filter(role=User.ROLE_AHC):
            send_notification(ahc_user, 'repair', f'Новая заявка {req.number}', f'Кабинет: {room}. {description[:100]}', f'/requests/{req.id}/')
        messages.success(request, f'Заявка {req.number} создана')
        return redirect('request_detail', pk=req.pk)
    equipment_by_room = {}
    for room in rooms:
        from equipment.models import Equipment as Eq
        equipment_by_room[room.id] = list(Eq.objects.filter(room=room).values('id', 'brand', 'model', 'reg_number'))
    import json
    return render(request, 'requests_app/create.html', {
        'rooms': rooms,
        'priorities': RepairRequest.PRIORITY_CHOICES,
        'equipment_json': json.dumps(equipment_by_room),
    })


@login_required
def request_update_status(request, pk):
    req = get_object_or_404(RepairRequest, pk=pk)
    if request.method == 'POST':
        new_status = request.POST.get('status')
        ahc_response = request.POST.get('ahc_response', '')
        assigned_to_id = request.POST.get('assigned_to')
        if new_status in dict(RepairRequest.STATUS_CHOICES):
            req.status = new_status
        if ahc_response:
            req.ahc_response = ahc_response
        if assigned_to_id:
            req.assigned_to = User.objects.filter(pk=assigned_to_id).first()
        req.save()
        # Notify creator
        if req.created_by:
            send_notification(req.created_by, 'answer', f'Обновление заявки {req.number}',
                f'Статус изменён на: {req.get_status_display()}. {ahc_response}', f'/requests/{req.id}/')
        messages.success(request, 'Статус обновлён')
    return redirect('request_detail', pk=pk)


@login_required
def request_download_docx(request, pk):
    req = get_object_or_404(RepairRequest, pk=pk)
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        doc.add_heading('ЗАЯВКА НА РЕМОНТ ОБОРУДОВАНИЯ', 0)
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        rows_data = [
            ('Номер заявки', req.number),
            ('Кабинет', str(req.room)),
            ('Оборудование', str(req.equipment) if req.equipment else '—'),
            ('Описание', req.description),
            ('Приоритет', req.get_priority_display()),
            ('Статус', req.get_status_display()),
            ('Заявитель', req.created_by.get_full_name() if req.created_by else '—'),
            ('Дата создания', req.created_at.strftime('%d.%m.%Y %H:%M')),
        ]
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="request_{req.number}.docx"'
        return response
    except ImportError:
        return HttpResponse('python-docx не установлен', status=500)


@login_required
def change_request_list(request):
    if request.user.is_admin:
        qs = ChangeRequest.objects.select_related('room', 'created_by').all()
    else:
        qs = ChangeRequest.objects.filter(created_by=request.user)
    return render(request, 'requests_app/change_list.html', {'requests': qs})


@login_required
def change_request_create(request):
    rooms = request.user.rooms_kept.all() if request.user.is_keeper else Room.objects.all()
    if request.method == 'POST':
        room = get_object_or_404(Room, pk=request.POST.get('room'))
        req = ChangeRequest.objects.create(
            room=room,
            request_type=request.POST.get('request_type'),
            description=request.POST.get('description'),
            created_by=request.user,
        )
        for admin in User.objects.filter(role=User.ROLE_ADMIN):
            send_notification(admin, 'approval', f'Запрос на изменение кабинета {room}',
                f'От: {request.user}. {req.description[:100]}', '/requests/change-requests/')
        messages.success(request, 'Запрос отправлен администратору')
        return redirect('change_request_list')
    return render(request, 'requests_app/change_create.html', {
        'rooms': rooms,
        'types': ChangeRequest.TYPE_CHOICES,
    })


@login_required
def change_request_review(request, pk):
    if not request.user.is_admin:
        return redirect('dashboard')
    req = get_object_or_404(ChangeRequest, pk=pk)
    if request.method == 'POST':
        req.status = request.POST.get('status')
        req.admin_comment = request.POST.get('admin_comment', '')
        req.reviewed_by = request.user
        req.reviewed_at = timezone.now()
        req.save()
        if req.created_by:
            send_notification(req.created_by, 'approval', f'Ответ на запрос: {req.room}',
                f'Статус: {req.get_status_display()}. {req.admin_comment}')
        messages.success(request, 'Запрос обработан')
        return redirect('change_request_list')
    return render(request, 'requests_app/change_review.html', {'req': req})
