from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
from django.db.models import Sum, Q
import openpyxl
from openpyxl.styles import Font, PatternFill
from equipment.models import Equipment, WriteOffRequest
from core.models import Floor


@login_required
def accounting_view(request):
    if not (request.user.is_admin or request.user.is_accountant):
        return redirect('dashboard')
    qs = Equipment.objects.select_related('room__floor__building', 'equipment_type')
    floor_id = request.GET.get('floor')
    room_id = request.GET.get('room')
    condition = request.GET.get('condition')
    if floor_id:
        qs = qs.filter(room__floor_id=floor_id)
    if room_id:
        qs = qs.filter(room_id=room_id)
    if condition:
        qs = qs.filter(condition=condition)

    total_cost = qs.filter(cost__isnull=False).aggregate(s=Sum('cost'))['s'] or 0
    active_cost = qs.filter(cost__isnull=False).exclude(condition='written_off').aggregate(s=Sum('cost'))['s'] or 0
    written_off_cost = qs.filter(cost__isnull=False, condition='written_off').aggregate(s=Sum('cost'))['s'] or 0

    from core.models import Room
    ctx = {
        'equipment': qs,
        'floors': Floor.objects.select_related('building').all(),
        'rooms': Room.objects.all(),
        'conditions': Equipment.CONDITION_CHOICES,
        'total_cost': total_cost,
        'active_cost': active_cost,
        'written_off_cost': written_off_cost,
        'filters': {'floor': floor_id, 'room': room_id, 'condition': condition},
    }
    return render(request, 'accounting/view.html', ctx)


@login_required
def update_cost(request):
    if not (request.user.is_admin or request.user.is_accountant):
        return redirect('dashboard')
    if request.method == 'POST':
        eq_id = request.POST.get('equipment_id')
        cost = request.POST.get('cost')
        Equipment.objects.filter(pk=eq_id).update(cost=cost or None)
        from core.models import ChangeLog
        ChangeLog.objects.create(user=request.user, action=f'Изменена стоимость оборудования ID={eq_id} → {cost}', model_name='Equipment', object_id=eq_id)
    return redirect('accounting')


@login_required
def accounting_export(request):
    if not (request.user.is_admin or request.user.is_accountant):
        return redirect('dashboard')
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Баланс оборудования'
    headers = ['Рег. номер', 'Марка', 'Модель', 'Кабинет', 'Этаж', 'Состояние', 'Количество', 'Стоимость (руб.)', 'Дата внесения']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill('solid', fgColor='1F3864')
    for item in Equipment.objects.select_related('room__floor').all():
        ws.append([
            item.reg_number, item.brand, item.model,
            item.room.number if item.room else '',
            str(item.room.floor) if item.room else '',
            item.get_condition_display(),
            item.quantity,
            float(item.cost) if item.cost else '',
            str(item.date_added),
        ])
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="accounting_report.xlsx"'
    wb.save(response)
    return response


@login_required
def writeoff_list(request):
    if not (request.user.is_admin or request.user.is_accountant):
        return redirect('dashboard')
    writeoffs = WriteOffRequest.objects.select_related('equipment', 'created_by').order_by('-created_at')
    return render(request, 'accounting/writeoffs.html', {'writeoffs': writeoffs})


@login_required
def writeoff_export(request):
    if not (request.user.is_admin or request.user.is_accountant):
        return redirect('dashboard')
    import io
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('АКТ СПИСАНИЯ ОБОРУДОВАНИЯ', 0)
        writeoffs = WriteOffRequest.objects.select_related('equipment', 'created_by').order_by('-created_at')
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Рег. номер'
        hdr[1].text = 'Оборудование'
        hdr[2].text = 'Причина'
        hdr[3].text = 'Дата'
        for wo in writeoffs:
            row = table.add_row().cells
            row[0].text = wo.equipment.reg_number
            row[1].text = f'{wo.equipment.brand} {wo.equipment.model}'
            row[2].text = wo.reason
            row[3].text = wo.created_at.strftime('%d.%m.%Y')
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        response = HttpResponse(buf.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="writeoff_act.docx"'
        return response
    except ImportError:
        return HttpResponse('python-docx не установлен', status=500)


@login_required
def export_1c(request):
    if not (request.user.is_admin or request.user.is_accountant):
        return redirect('dashboard')
    lines = ['НомерОС;Наименование;Количество;Стоимость;Состояние;Кабинет']
    for item in Equipment.objects.select_related('room').all():
        lines.append(f'{item.reg_number};{item.brand} {item.model};{item.quantity};{item.cost or 0};{item.get_condition_display()};{item.room.number if item.room else ""}')
    content = '\n'.join(lines)
    response = HttpResponse(content, content_type='text/csv; charset=utf-8-sig')
    response['Content-Disposition'] = 'attachment; filename="1c_export.csv"'
    return response
